"""
docx_ingest.py

DOCX ingestion module for Day 16 (Document Ingestion & Chunking Strategies).

Strategy:
    - python-docx walks the document's XML tree (word/document.xml under the
      hood) as a sequence of Paragraph and Table objects, in document order.
    - Unlike PDF, DOCX has no real concept of "pages" — page breaks are a
      rendering-time detail, not stored positions — so page_number is always
      None here. That's expected and matches metadata_schema.py's Optional type.
    - Heading detection is NOT a heuristic here (unlike pdf_ingest.py's font
      size guessing) — Word paragraphs carry an explicit style name
      ("Heading 1", "Heading 2", ...) baked into the XML, so we just read
      paragraph.style.name directly. This is one of the concrete advantages
      of DOCX's structured format over PDF's positional one.
    - Tables are routed through table_registry.py (same pattern as
      pdf_ingest.py): never inlined into the main element stream. Each
      table gets a table_id, an LLM (Gemini) summary, and an index entry in
      tables_lookup.json / tables_summary.md — the element left behind here
      is just a short pointer, so chunking never sees raw table content.

Output shape:
    extract_docx(path) -> list[ExtractedElement]
    One element per paragraph (heading or body) and one pointer element per
    table, in document order, each tagged with the nearest preceding heading.
"""

import sys
from pathlib import Path
from typing import Optional

# See pdf_ingest.py for why this is needed — makes metadata_schema
# importable regardless of the working directory this script is run from.
sys.path.append(str(Path(__file__).resolve().parent.parent))

import docx  # python-docx; package name differs from import name

from metadata_schema import ExtractedElement, ElementType, elements_to_json
from table_registry import (
    detect_caption, generate_table_id, summarize_table_with_llm, register_table,
)


def _is_heading_style(style_name: str) -> bool:
    """
    Word's built-in heading styles are literally named "Heading 1",
    "Heading 2", etc. Custom templates sometimes rename them, but this
    covers the vast majority of real-world .docx files.
    """
    return style_name.lower().startswith("heading")


def _is_list_style(paragraph: "docx.text.paragraph.Paragraph") -> bool:
    """
    Detects bullet/numbered list items. python-docx doesn't expose a clean
    'is_list' flag, so we check two things:
      1. Direct numPr on the paragraph itself (numbering applied ad-hoc).
      2. The paragraph's STYLE name containing "list" (e.g. Word's built-in
         "List Bullet" / "List Number" styles) — when you apply one of these
         built-in styles, the numbering is often inherited from the style
         definition rather than stamped onto the paragraph's own XML, so
         numPr alone misses it.
    """
    pPr = paragraph._p.pPr
    has_direct_numpr = pPr is not None and pPr.numPr is not None

    style_name = paragraph.style.name if paragraph.style else ""
    has_list_style = "list" in style_name.lower()

    return has_direct_numpr or has_list_style


def _table_to_text(table: "docx.table.Table") -> str:
    """Serializes a python-docx Table into the same pipe-delimited format
    used by pdf_ingest.py, so downstream summarization treats both
    consistently."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(" | ".join(cells))
    return "\n".join(rows)


def extract_docx(path: str) -> list[ExtractedElement]:
    """
    Extracts a DOCX into an ordered list of ExtractedElement.

    Walks doc.element.body children directly (rather than doc.paragraphs and
    doc.tables separately) so paragraphs and tables come out in their TRUE
    interleaved document order — iterating the two lists separately would
    lose that ordering whenever a table sits between two paragraphs. This
    order matters here specifically because caption detection relies on
    knowing which paragraph immediately preceded a given table.
    """
    source_filename = Path(path).name
    document = docx.Document(path)

    elements: list[ExtractedElement] = []
    element_index = 0
    current_heading: Optional[str] = None
    last_paragraph_text: Optional[str] = None
    table_index = 0

    # Build lookup so we can walk body children (in true order) and still
    # get python-docx's richer Paragraph/Table wrapper objects for each.
    paragraphs_by_xml = {p._p: p for p in document.paragraphs}
    tables_by_xml = {t._tbl: t for t in document.tables}

    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):  # w:p -> paragraph
            paragraph = paragraphs_by_xml.get(child)
            if paragraph is None:
                continue
            text = paragraph.text.strip()
            if not text:
                continue  # skip blank paragraphs (common as visual spacing in Word)

            style_name = paragraph.style.name if paragraph.style else ""

            if _is_heading_style(style_name):
                current_heading = text
                elements.append(ExtractedElement(
                    text=text,
                    source_filename=source_filename,
                    page_number=None,  # DOCX has no stored page positions
                    section_heading=current_heading,
                    element_type=ElementType.HEADING,
                    element_index=element_index,
                ))
            elif _is_list_style(paragraph):
                elements.append(ExtractedElement(
                    text=text,
                    source_filename=source_filename,
                    page_number=None,
                    section_heading=current_heading,
                    element_type=ElementType.LIST_ITEM,
                    element_index=element_index,
                ))
            else:
                elements.append(ExtractedElement(
                    text=text,
                    source_filename=source_filename,
                    page_number=None,
                    section_heading=current_heading,
                    element_type=ElementType.PARAGRAPH,
                    element_index=element_index,
                ))
            element_index += 1
            last_paragraph_text = text  # remember for caption detection on the next table

        elif child.tag.endswith("}tbl"):  # w:tbl -> table
            table = tables_by_xml.get(child)
            if table is None:
                continue
            raw_markdown = _table_to_text(table)
            if not raw_markdown.strip():
                continue  # skip genuinely empty tables

            table_id = generate_table_id(source_filename, table_index)
            table_index += 1

            # captions conventionally sit in the paragraph right before the
            # table (e.g. "Table 2.13: ..."), which is exactly what
            # last_paragraph_text is tracking above
            caption = detect_caption(last_paragraph_text)
            summary = summarize_table_with_llm(raw_markdown)

            register_table(
                table_id=table_id,
                source_filename=source_filename,
                page_number=None,
                section_heading=current_heading,
                caption=caption,
                raw_markdown=raw_markdown,
                summary=summary,
            )

            placeholder = f"[{caption or 'Table'} — see table registry, id: {table_id}]"
            elements.append(ExtractedElement(
                text=placeholder,
                source_filename=source_filename,
                page_number=None,
                section_heading=current_heading,
                element_type=ElementType.TABLE,
                element_index=element_index,
                table_id=table_id,
            ))
            element_index += 1

    return elements


# ---------------------------------------------------------------------------
# Script entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    from table_registry import TABLES_LOOKUP_PATH, TABLES_SUMMARY_PATH

    # NOTE: does NOT call reset_registry() here. If you're ingesting PDF,
    # DOCX, and OCR documents in one full pipeline run, reset the registry
    # ONCE at the very start of that run (before the first document), not
    # inside every individual ingestion script — otherwise each script
    # wipes out the tables the previous one just registered.

    RAW_PATH = sys.argv[1] if len(sys.argv) > 1 else "data/raw/sample.docx"
    OUT_PATH = "data/extracted/docx_text.json"

    elements = extract_docx(RAW_PATH)
    headings = [e for e in elements if e.element_type.value == "heading"]
    paragraphs = [e for e in elements if e.element_type.value == "paragraph"]
    list_items = [e for e in elements if e.element_type.value == "list_item"]
    tables = [e for e in elements if e.element_type.value == "table"]

    print(f"[docx_ingest] Extracted {len(elements)} elements from {RAW_PATH}")
    print(f"  {len(headings)} headings")
    print(f"  {len(paragraphs)} paragraphs")
    print(f"  {len(list_items)} list items")
    print(f"  {len(tables)} table pointers (full content routed to registry)")

    elements_to_json(elements, OUT_PATH)
    print(f"[docx_ingest] Wrote {OUT_PATH}")
    if tables:
        print(f"Table registry: {TABLES_LOOKUP_PATH}, {TABLES_SUMMARY_PATH}")